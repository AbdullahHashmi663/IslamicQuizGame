import os
import django
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up Django environment
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'islamic_quiz.settings')
django.setup()

from quiz_app.models import Quiz, Question, Answer

def create_quiz_solutions_doc():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Islamic Quiz Solutions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Get all quizzes
    quizzes = Quiz.objects.all()
    
    for quiz in quizzes:
        # Add quiz title
        doc.add_heading(quiz.title, level=1)
        doc.add_paragraph(quiz.description)
        
        # Get all questions for this quiz
        questions = Question.objects.filter(quiz=quiz)
        
        for i, question in enumerate(questions, 1):
            # Add question
            q_para = doc.add_paragraph()
            q_para.add_run(f'Question {i}: ').bold = True
            q_para.add_run(question.text)
            
            # Add options
            answers = Answer.objects.filter(question=question)
            for answer in answers:
                option_para = doc.add_paragraph()
                option_para.style = 'List Bullet'
                option_para.add_run(answer.text)
                if answer.is_correct:
                    option_para.add_run(' (Correct Answer)').bold = True
            
            # Add spacing between questions
            doc.add_paragraph()
    
    # Save the document
    doc.save('Quiz_Solutions.docx')

if __name__ == '__main__':
    create_quiz_solutions_doc() 